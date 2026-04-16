import os
import re
import json
import requests
import gradio as gr
import time
import fitz
import docx
from dotenv import load_dotenv
from collections import deque, Counter
from threading import Lock
import tempfile
import html as html_lib
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
load_dotenv()

# ================== SUPPORTED LANGUAGES ==================
SUPPORTED_LANGUAGES = [
    "Afrikaans","Albanian","Amharic","Arabic","Armenian","Assamese","Azerbaijani",
    "Basque","Belarusian","Bengali","Bosnian","Bulgarian","Burmese",
    "Catalan","Cebuano","Chinese (Simplified)","Chinese (Traditional)","Corsican","Croatian","Czech",
    "Danish","Dutch","English","Esperanto","Estonian",
    "Filipino","Finnish","French","Frisian",
    "Galician","Georgian","German","Greek","Gujarati",
    "Haitian Creole","Hausa","Hawaiian","Hebrew","Hindi","Hmong","Hungarian",
    "Icelandic","Igbo","Indonesian","Irish","Italian",
    "Japanese","Javanese",
    "Kannada","Kazakh","Khmer","Kinyarwanda","Korean","Kurdish","Kyrgyz",
    "Lao","Latin","Latvian","Lithuanian","Luxembourgish",
    "Macedonian","Malagasy","Malay","Malayalam","Maltese","Maori","Marathi","Mongolian",
    "Nepali","Norwegian",
    "Odia","Pashto","Persian","Polish","Portuguese","Punjabi",
    "Romanian","Russian",
    "Samoan","Scots Gaelic","Serbian","Sesotho","Shona","Sindhi","Sinhala","Slovak","Slovenian",
    "Somali","Spanish","Sundanese","Swahili","Swedish",
    "Tagalog","Tajik","Tamil","Tatar","Telugu","Thai","Turkish","Turkmen",
    "Ukrainian","Urdu","Uyghur","Uzbek",
    "Vietnamese","Welsh","Xhosa","Yiddish","Yoruba","Zulu"
]

# ================== CONFIG ==================
API_URL  = "https://openrouter.ai/api/v1/chat/completions"
API_KEY  = os.environ.get("OPENROUTER_API_KEY", "")
MODEL    = "arcee-ai/trinity-large-preview:free"

CHUNK_SIZE_WORDS     = 250
MAX_TOKENS_PER_CHUNK = 1024
CHUNKS_PER_BATCH     = 8
RATE_WAIT_SECS       = 5
MAX_RETRIES          = 6
RETRY_DELAYS         = [5, 10, 20, 40, 60, 90]

MAX_REQUESTS_PER_MINUTE = 20
MAX_REQUESTS_PER_DAY    = 500
MINUTE_WINDOW = 60
DAY_WINDOW    = 24 * 60 * 60
minute_timestamps = deque()
day_timestamps    = deque()
rate_limit_lock   = Lock()

# ================== LANGUAGE-TONE VOCABULARY ENGINE ==================
# This will be injected into app.py
# Complete language + tone vocabulary seeding for all supported languages

# ── Script family detection ────────────────────────────────────────────────────
INDIC_LANGS = {
    "Assamese","Bengali","Gujarati","Hindi","Kannada","Malayalam","Marathi",
    "Nepali","Odia","Punjabi","Sinhala","Tamil","Telugu","Urdu","Sanskrit"
}
ARABIC_SCRIPT_LANGS = {"Arabic","Persian","Pashto","Sindhi","Uyghur","Urdu"}
CJK_LANGS = {"Chinese (Simplified)","Chinese (Traditional)","Japanese","Korean"}
CYRILLIC_LANGS = {
    "Belarusian","Bulgarian","Kazakh","Kyrgyz","Macedonian",
    "Mongolian","Russian","Serbian","Tajik","Tatar","Ukrainian"
}
RTL_LANGS = ARABIC_SCRIPT_LANGS | {"Hebrew","Yiddish"}

def get_script_family(lang):
    if lang in INDIC_LANGS:      return "indic"
    if lang in ARABIC_SCRIPT_LANGS: return "arabic"
    if lang in CJK_LANGS:        return "cjk"
    if lang in CYRILLIC_LANGS:   return "cyrillic"
    if lang in RTL_LANGS:        return "rtl"
    return "latin"

# ── Tone vocabulary seeds for major non-Latin languages ───────────────────────
# Format: lang → tone → (emotional_words_in_that_language, writing_style_note)
LANG_TONE_VOCAB = {
    "Telugu": {
        "happy":        ("సంతోషం, ఆనందం, హర్షం, ఉల్లాసం",       "వాక్యాలు చిన్నవిగా, ఉత్సాహంగా ఉండాలి"),
        "sad":          ("దుఃఖం, విషాదం, బాధ, నిరాశ",             "నెమ్మదిగా, భావయుక్తంగా రాయాలి"),
        "romantic":     ("ప్రేమ, అనురాగం, ఆకర్షణ, హృదయం",        "మృదువుగా, కవిత్వంగా రాయాలి"),
        "angry":        ("కోపం, ఆగ్రహం, రోషం, ఆవేశం",             "తీవ్రంగా, నేరుగా రాయాలి"),
        "poetic":       ("కవిత్వం, భావం, చిత్రణ, లయ",             "అలంకారాలు, రూపకాలు వాడాలి"),
        "motivational": ("స్ఫూర్తి, ధైర్యం, శక్తి, పురోగతి",      "శక్తిమంతంగా, ప్రోత్సాహంగా రాయాలి"),
        "formal":       ("గౌరవం, నిష్పక్షపాతం, అధికారిక భాష",    "వ్యాకరణబద్ధంగా, అధికారికంగా రాయాలి"),
        "calm":         ("శాంతి, ప్రశాంతత, నిశ్చలత",               "నెమ్మదిగా, హాయిగా రాయాలి"),
        "fearful":      ("భయం, ఆందోళన, గాభరా",                   "అసహజంగా, తడబాటుతో రాయాలి"),
        "nostalgic":    ("గతం, జ్ఞాపకాలు, మిస్సింగ్",             "మధురంగా, గతాన్ని తలచుకుంటూ రాయాలి"),
    },
    "Hindi": {
        "happy":        ("खुशी, आनंद, प्रसन्नता, हर्ष",           "छोटे उत्साही वाक्य लिखें"),
        "sad":          ("दुख, उदासी, विषाद, पीड़ा",               "धीमे, भावपूर्ण वाक्य लिखें"),
        "romantic":     ("प्यार, मोहब्बत, इश्क, दिल",             "कोमल, काव्यात्मक भाषा लिखें"),
        "angry":        ("क्रोध, गुस्सा, रोष, आक्रोश",             "तीव्र, सीधे वाक्य लिखें"),
        "poetic":       ("काव्य, भाव, छंद, अलंकार",               "रूपकों और बिम्बों का प्रयोग करें"),
        "motivational": ("प्रेरणा, साहस, शक्ति, उत्साह",          "शक्तिशाली, प्रोत्साहित करने वाली भाषा"),
        "formal":       ("औपचारिक, सम्मानजनक, शिष्ट",             "मानक हिंदी में लिखें"),
        "calm":         ("शांति, सुकून, स्थिरता",                  "धीमी, सुखद लय में लिखें"),
        "fearful":      ("डर, भय, आशंका, घबराहट",                 "टूटे वाक्य, अनिश्चितता दर्शाएं"),
        "nostalgic":    ("यादें, पुराना, बचपन, अतीत",              "मीठी यादों के साथ लिखें"),
    },
    "Tamil": {
        "happy":        ("மகிழ்ச்சி, ஆனந்தம், சந்தோஷம், உவகை",  "குறுகிய, உற்சாகமான வாக்கியங்கள்"),
        "sad":          ("துக்கம், சோகம், வேதனை, கவலை",           "மெதுவான, உணர்ச்சிமிக்க வாக்கியங்கள்"),
        "romantic":     ("காதல், அன்பு, பாசம், நேசம்",            "மென்மையான, கவித்துவமான மொழி"),
        "angry":        ("கோபம், சினம், ஆத்திரம், வெகுண்டு",     "தீவிரமான, நேரடியான வாக்கியங்கள்"),
        "poetic":       ("கவிதை, இலக்கியம், உருவகம், தாளம்",      "அலங்காரங்கள் மற்றும் உருவகங்கள்"),
        "motivational": ("உத்வேகம், தைரியம், சக்தி, முன்னேற்றம்","சக்திவாய்ந்த, ஊக்கமளிக்கும் மொழி"),
        "calm":         ("அமைதி, நிதானம், சாந்தி",                "மெல்லிய, இனிமையான தாளம்"),
        "formal":       ("அதிகாரப்பூர்வம், மரியாதை, சட்டப்பூர்வம்","தமிழ் இலக்கண விதிகளை பின்பற்றவும்"),
    },
    "Bengali": {
        "happy":        ("আনন্দ, সুখ, খুশি, হর্ষ",               "ছোট উৎসাহী বাক্য লিখুন"),
        "sad":          ("দুঃখ, বেদনা, কষ্ট, বিষাদ",              "ধীর, আবেগপূর্ণ বাক্য লিখুন"),
        "romantic":     ("ভালোবাসা, প্রেম, রোমান্স, হৃদয়",       "কোমল, কাব্যিক ভাষা ব্যবহার করুন"),
        "angry":        ("রাগ, ক্রোধ, ক্ষোভ, উত্তেজনা",          "তীব্র, সরাসরি বাক্য লিখুন"),
        "poetic":       ("কবিতা, ছন্দ, রূপক, চিত্রকল্প",          "রবীন্দ্রনাথের মতো ভাষা"),
        "nostalgic":    ("স্মৃতি, অতীত, শৈশব, মিস",              "মিষ্টি স্মৃতি দিয়ে লিখুন"),
        "calm":         ("শান্তি, প্রশান্তি, স্থিরতা",             "ধীর, সুখকর তালে লিখুন"),
        "formal":       ("আনুষ্ঠানিক, সম্মানজনক, ভদ্র",          "প্রমিত বাংলায় লিখুন"),
    },
    "Malayalam": {
        "happy":        ("സന്തോഷം, ആനന്ദം, സുഖം, ഹർഷം",          "ചെറിയ ഉത്സാഹഭരിത വാക്യങ്ങൾ"),
        "sad":          ("സങ്കടം, ദുഃഖം, വ്യസനം, നൊമ്പരം",        "സാവധാനം, വൈകാരിക വാക്യങ്ങൾ"),
        "romantic":     ("പ്രണയം, സ്നേഹം, ആകർഷണം, ഹൃദയം",       "മൃദുവായ, കാവ്യാത്മക ഭാഷ"),
        "angry":        ("കോപം, ദേഷ്യം, രോഷം, ക്ഷോഭം",           "തീവ്രമായ, നേരിട്ടുള്ള വാക്യങ്ങൾ"),
        "calm":         ("ശാന്തത, സമാധാനം, നിശ്ചലത",             "ശാന്തവും സുഖകരവുമായ താളം"),
        "poetic":       ("കവിത, ഭാവം, ഉപമ, ലയം",                  "അലങ്കാരങ്ങളും രൂപകങ്ങളും"),
    },
    "Kannada": {
        "happy":        ("ಸಂತೋಷ, ಆನಂದ, ಹರ್ಷ, ಉಲ್ಲಾಸ",           "ಚಿಕ್ಕ ಉತ್ಸಾಹದ ವಾಕ್ಯಗಳು"),
        "sad":          ("ದುಃಖ, ವಿಷಾದ, ನೋವು, ಬೇಸರ",              "ನಿಧಾನ, ಭಾವಪೂರ್ಣ ವಾಕ್ಯಗಳು"),
        "romantic":     ("ಪ್ರೇಮ, ಅನುರಾಗ, ಆಕರ್ಷಣ, ಹೃದಯ",          "ಮೃದು, ಕಾವ್ಯಾತ್ಮಕ ಭಾಷೆ"),
        "angry":        ("ಕೋಪ, ಕ್ರೋಧ, ರೋಷ, ಆವೇಶ",               "ತೀವ್ರ, ನೇರ ವಾಕ್ಯಗಳು"),
        "calm":         ("ಶಾಂತಿ, ಪ್ರಶಾಂತತೆ, ನಿರ್ಮಲ",             "ನಿಧಾನ, ಸುಖಕರ ಲಯ"),
    },
    "Gujarati": {
        "happy":        ("ખુશી, આનંદ, પ્રસન્નતા, હર્ષ",           "ટૂંકા ઉત્સાહી વાક્યો"),
        "sad":          ("દુઃખ, ઉદાસી, વ્યથા, ખેદ",               "ધીમા, ભાવભર્યા વાક્યો"),
        "romantic":     ("પ્રેમ, ચાહત, લાગણી, દિલ",              "કોમળ, કાવ્યાત્મક ભાષા"),
        "angry":        ("ગુસ્સો, ક્રોધ, રોષ, આક્રોશ",            "તીવ્ર, સીધા વાક્યો"),
        "calm":         ("શાંતિ, સ્થિરતા, સુકૂન",                  "ધીમી, સુખદ લય"),
    },
    "Punjabi": {
        "happy":        ("ਖੁਸ਼ੀ, ਆਨੰਦ, ਹਰਸ਼, ਖੇੜਾ",              "ਛੋਟੇ ਉਤਸ਼ਾਹੀ ਵਾਕ"),
        "sad":          ("ਦੁੱਖ, ਉਦਾਸੀ, ਗ਼ਮ, ਦਰਦ",                "ਹੌਲੀ, ਭਾਵੁਕ ਵਾਕ"),
        "romantic":     ("ਪਿਆਰ, ਮੁਹੱਬਤ, ਇਸ਼ਕ, ਦਿਲ",             "ਕੋਮਲ, ਕਾਵਿਕ ਭਾਸ਼ਾ"),
        "angry":        ("ਗੁੱਸਾ, ਕ੍ਰੋਧ, ਰੋਹ, ਆਕਰੋਸ਼",            "ਤਿੱਖੇ, ਸਿੱਧੇ ਵਾਕ"),
        "calm":         ("ਸ਼ਾਂਤੀ, ਸਕੂਨ, ਠਹਿਰਾਅ",                  "ਹੌਲੀ, ਸੁਖਦ ਲੈਅ"),
    },
    "Marathi": {
        "happy":        ("आनंद, सुख, प्रसन्नता, हर्ष",             "लहान उत्साही वाक्ये"),
        "sad":          ("दुःख, विषाद, वेदना, खेद",                "हळू, भावपूर्ण वाक्ये"),
        "romantic":     ("प्रेम, माया, आकर्षण, मन",               "कोमल, काव्यात्मक भाषा"),
        "angry":        ("राग, क्रोध, रोष, संताप",                 "तीव्र, थेट वाक्ये"),
        "calm":         ("शांती, स्थिरता, निर्मळता",               "हळू, सुखद लय"),
        "formal":       ("औपचारिक, आदरपूर्वक, शिष्ट",             "मानक मराठीत लिहा"),
    },
    "Nepali": {
        "happy":        ("खुशी, आनन्द, प्रसन्नता, हर्ष",           "छोटो उत्साही वाक्यहरू"),
        "sad":          ("दुःख, उदासी, पीडा, विषाद",               "ढिलो, भावपूर्ण वाक्यहरू"),
        "romantic":     ("प्रेम, माया, स्नेह, दिल",               "कोमल, काव्यात्मक भाषा"),
        "angry":        ("रिस, क्रोध, रोष, आक्रोश",               "तीव्र, प्रत्यक्ष वाक्यहरू"),
        "calm":         ("शान्ति, स्थिरता, सुकून",                  "ढिलो, सुखद लय"),
    },
    "Sinhala": {
        "happy":        ("සතුට, ආනන්දය, සුභ, ප්‍රීතිය",           "කෙටි උද්‍යෝගිමත් වාක්‍ය"),
        "sad":          ("දුක, වේදනාව, කනස්සල්ල",                 "සෙමින්, හැඟීම්බර වාක්‍ය"),
        "romantic":     ("ආදරය, හදවත, ප්‍රේමය",                   "මෘදු, කාව්‍යාත්මක භාෂාව"),
        "angry":        ("කෝපය, රෝෂය, තරහ",                      "තීව්‍ර, සෘජු වාක්‍ය"),
        "calm":         ("සාමය, නිදහස, ස්ථිරතාව",                 "සෙමින්, සුඛදායක රිද්මය"),
    },
    "Odia": {
        "happy":        ("ଆନନ୍ଦ, ସୁଖ, ହର୍ଷ, ଖୁସି",               "ଛୋଟ ଉତ୍ସାହ ବାକ୍ୟ"),
        "sad":          ("ଦୁଃଖ, ବ୍ୟଥା, ବିଷାଦ, ଯନ୍ତ୍ରଣା",          "ଧୀର, ଭାବ ବାକ୍ୟ"),
        "romantic":     ("ପ୍ରେମ, ଭଲଲାଗ, ଆକର୍ଷଣ",                "ମୃଦୁ, କାବ୍ୟ ଭାଷା"),
        "calm":         ("ଶାନ୍ତ, ସ୍ଥିରତା, ନିର୍ମଳ",               "ଧୀର, ସୁଖଦ ତାଳ"),
    },
    "Arabic": {
        "happy":        ("سعادة، فرح، بهجة، سرور",                "جمل قصيرة ومبهجة"),
        "sad":          ("حزن، أسى، غم، كآبة",                   "جمل بطيئة وعاطفية"),
        "romantic":     ("حب، عشق، غرام، رومانسية",              "لغة ناعمة وشاعرية"),
        "angry":        ("غضب، سخط، حنق، ثورة",                  "جمل حادة ومباشرة"),
        "poetic":       ("شعر، قصيدة، استعارة، وزن",              "استخدم الصور الشعرية"),
        "motivational": ("إلهام، شجاعة، قوة، تحفيز",             "لغة قوية وملهمة"),
        "formal":       ("رسمي، محترم، مهني",                    "استخدم الفصحى الرسمية"),
        "calm":         ("هدوء، سكينة، طمأنينة",                  "إيقاع هادئ وسلس"),
        "fearful":      ("خوف، رهبة، قلق، توتر",                 "جمل متقطعة ومترددة"),
        "nostalgic":    ("ذكريات، ماضي، شوق، حنين",              "كتابة حنينية دافئة"),
    },
    "Persian": {
        "happy":        ("شادی، خوشحالی، نشاط، سرور",            "جملات کوتاه و پرانرژی"),
        "sad":          ("غم، اندوه، دلتنگی، حسرت",              "جملات آرام و احساساتی"),
        "romantic":     ("عشق، محبت، دلدادگی، قلب",              "زبان لطیف و شاعرانه"),
        "angry":        ("خشم، غضب، عصبانیت، برافروختگی",        "جملات تند و مستقیم"),
        "poetic":       ("شعر، غزل، استعاره، وزن",                "از تصاویر شاعرانه استفاده کن"),
        "calm":         ("آرامش، سکوت، طمانینه",                  "ریتم آرام و دلنشین"),
        "nostalgic":    ("خاطره، گذشته، دلتنگی، نوستالژی",       "نوشتن با حس دلتنگی گرم"),
    },
    "Urdu": {
        "happy":        ("خوشی، مسرت، شادمانی، انبساط",           "چھوٹے پرجوش جملے"),
        "sad":          ("غم، اداسی، دکھ، ملال",                  "آہستہ، جذباتی جملے"),
        "romantic":     ("محبت، عشق، الفت، دل",                  "نرم، شاعرانہ زبان"),
        "angry":        ("غصہ، غضب، قہر، غیظ",                   "تیز، سیدھے جملے"),
        "poetic":       ("شاعری، غزل، استعارہ، وزن",              "میر، غالب کی طرح لکھیں"),
        "formal":       ("رسمی، با ادب، پیشہ ورانہ",              "فصیح اردو میں لکھیں"),
        "calm":         ("سکون، اطمینان، قرار",                   "آہستہ، سکون بخش لہجہ"),
        "nostalgic":    ("یادیں، ماضی، بچپن، یاد",               "گرم یادوں کے ساتھ لکھیں"),
    },
    "Chinese (Simplified)": {
        "happy":        ("快乐、幸福、喜悦、欢乐",                  "简短有活力的句子"),
        "sad":          ("悲伤、忧郁、痛苦、难过",                  "缓慢、情感丰富的句子"),
        "romantic":     ("爱情、浪漫、温柔、心动",                  "柔和、诗意的语言"),
        "angry":        ("愤怒、怒火、激动、怒气",                  "直接、有力的句子"),
        "poetic":       ("诗意、意境、比喻、韵律",                  "使用中国古典意象"),
        "calm":         ("平静、安宁、淡然、宁静",                  "缓慢、舒缓的节奏"),
        "formal":       ("正式、庄重、专业",                        "使用正式书面语"),
        "motivational": ("激励、勇气、力量、拼搏",                  "有力量的激励语言"),
    },
    "Chinese (Traditional)": {
        "happy":        ("快樂、幸福、喜悅、歡樂",                  "簡短有活力的句子"),
        "sad":          ("悲傷、憂鬱、痛苦、難過",                  "緩慢、情感豐富的句子"),
        "romantic":     ("愛情、浪漫、溫柔、心動",                  "柔和、詩意的語言"),
        "angry":        ("憤怒、怒火、激動、怒氣",                  "直接、有力的句子"),
        "poetic":       ("詩意、意境、比喻、韻律",                  "使用中國古典意象"),
        "calm":         ("平靜、安寧、淡然、寧靜",                  "緩慢、舒緩的節奏"),
    },
    "Japanese": {
        "happy":        ("幸せ、喜び、楽しい、嬉しい",             "短く活気のある文を書く"),
        "sad":          ("悲しい、悲しみ、切ない、哀愁",           "ゆっくりと感情的な文を書く"),
        "romantic":     ("愛、恋、ロマンス、想い",                 "柔らかく詩的な言葉を使う"),
        "angry":        ("怒り、憤り、激怒、憤慨",                 "鋭く直接的な文を書く"),
        "poetic":       ("詩的、風情、比喩、情景",                  "季語や和歌のイメージを使う"),
        "calm":         ("静か、平和、穏やか、落ち着き",            "ゆっくりとした心地よいリズム"),
        "formal":       ("丁寧、礼儀、正式、敬語",                  "丁寧語・敬語を使う"),
        "nostalgic":    ("懐かしい、思い出、昔、あの頃",           "温かい思い出とともに書く"),
    },
    "Korean": {
        "happy":        ("행복, 기쁨, 즐거움, 환희",               "짧고 활기찬 문장을 쓰세요"),
        "sad":          ("슬픔, 悲, 고통, 애수",                   "느리고 감정적인 문장을 쓰세요"),
        "romantic":     ("사랑, 로맨스, 설렘, 마음",               "부드럽고 시적인 언어를 사용하세요"),
        "angry":        ("분노, 화, 격분, 울분",                   "직접적이고 강한 문장을 쓰세요"),
        "poetic":       ("시적, 운율, 은유, 서정",                 "한국 시의 서정적 이미지 사용"),
        "calm":         ("평온, 고요, 안정, 평화",                 "느리고 편안한 리듬"),
        "formal":       ("공식적, 격식, 예의, 정중",               "존댓말을 사용하세요"),
        "nostalgic":    ("그리움, 추억, 옛날, 향수",               "따뜻한 추억과 함께 쓰세요"),
    },
    "Russian": {
        "happy":        ("счастье, радость, веселье, восторг",    "Короткие энергичные предложения"),
        "sad":          ("грусть, печаль, тоска, горе",           "Медленные, эмоциональные предложения"),
        "romantic":     ("любовь, романтика, нежность, сердце",   "Мягкий, поэтический язык"),
        "angry":        ("злость, гнев, ярость, возмущение",      "Резкие, прямые предложения"),
        "poetic":       ("поэзия, метафора, образ, ритм",         "Использовать поэтические образы"),
        "formal":       ("официальный, уважительный, деловой",    "Использовать официальный стиль"),
        "nostalgic":    ("воспоминания, прошлое, ностальгия",     "Писать с теплой ностальгией"),
        "calm":         ("спокойствие, тишина, покой",            "Медленный, умиротворяющий ритм"),
    },
    "Ukrainian": {
        "happy":        ("щастя, радість, веселощі, захоплення",  "Короткі енергійні речення"),
        "sad":          ("смуток, печаль, туга, горе",            "Повільні, емоційні речення"),
        "romantic":     ("кохання, романтика, ніжність",          "М'яка, поетична мова"),
        "angry":        ("злість, гнів, лють, обурення",          "Різкі, прямі речення"),
        "calm":         ("спокій, тиша, умиротворення",           "Повільний, приємний ритм"),
        "nostalgic":    ("спогади, минуле, ностальгія",           "Писати з теплою ностальгією"),
    },
    "Hebrew": {
        "happy":        ("שמחה, אושר, ששון, עליזות",             "משפטים קצרים ואנרגטיים"),
        "sad":          ("עצב, צער, כאב, יגון",                  "משפטים איטיים ורגשיים"),
        "romantic":     ("אהבה, רומנטיקה, חיבה, לב",             "שפה עדינה ופיוטית"),
        "angry":        ("כעס, זעם, רוגז, חמה",                  "משפטים חדים וישירים"),
        "calm":         ("שלווה, רגיעה, שקט",                    "קצב איטי ונעים"),
        "formal":       ("רשמי, מנומס, מכובד",                   "שימוש בעברית גבוהה"),
    },
    "Greek": {
        "happy":        ("ευτυχία, χαρά, αγαλλίαση, αγαλλίαση", "Σύντομες ενεργητικές προτάσεις"),
        "sad":          ("λύπη, θλίψη, πόνος, μελαγχολία",       "Αργές, συναισθηματικές προτάσεις"),
        "romantic":     ("αγάπη, έρωτας, ρομαντισμός",           "Απαλή, ποιητική γλώσσα"),
        "angry":        ("θυμός, οργή, αγανάκτηση",              "Απότομες, άμεσες προτάσεις"),
        "calm":         ("ηρεμία, γαλήνη, ησυχία",               "Αργός, ευχάριστος ρυθμός"),
        "formal":       ("επίσημος, τυπικός, ευγενής",           "Χρησιμοποιήστε επίσημη ελληνική"),
    },
    "Turkish": {
        "happy":        ("mutluluk, neşe, sevinç, coşku",         "Kısa enerjik cümleler"),
        "sad":          ("üzüntü, hüzün, keder, acı",             "Yavaş, duygusal cümleler"),
        "romantic":     ("aşk, sevgi, romantizm, gönül",          "Yumuşak, şiirsel dil"),
        "angry":        ("öfke, kızgınlık, hiddet, sinir",        "Sert, doğrudan cümleler"),
        "poetic":       ("şiir, mecaz, imgelem, ritim",            "Türk şiir geleneğinden ilham al"),
        "calm":         ("huzur, sükunet, dinginlik",              "Yavaş, rahatlatıcı ritim"),
        "nostalgic":    ("nostalji, anılar, geçmiş, özlem",       "Sıcak anılarla yaz"),
    },
    "Vietnamese": {
        "happy":        ("hạnh phúc, vui vẻ, vui mừng, hân hoan","Câu ngắn, năng động"),
        "sad":          ("buồn, đau buồn, đau khổ, u sầu",        "Câu chậm, đầy cảm xúc"),
        "romantic":     ("tình yêu, lãng mạn, yêu thương",        "Ngôn ngữ nhẹ nhàng, thơ mộng"),
        "angry":        ("tức giận, giận dữ, phẫn nộ",            "Câu sắc bén, trực tiếp"),
        "calm":         ("bình yên, yên tĩnh, thanh thản",        "Nhịp điệu chậm, dễ chịu"),
        "formal":       ("trang trọng, lịch sự, chính thức",      "Dùng tiếng Việt chuẩn mực"),
    },
    "Indonesian": {
        "happy":        ("bahagia, gembira, senang, suka cita",   "Kalimat pendek dan energik"),
        "sad":          ("sedih, duka, nestapa, kesedihan",        "Kalimat lambat dan penuh perasaan"),
        "romantic":     ("cinta, kasih, romantis, hati",          "Bahasa lembut dan puitis"),
        "angry":        ("marah, amarah, murka, geram",            "Kalimat tajam dan langsung"),
        "calm":         ("tenang, damai, tentram",                 "Ritme lambat dan menyenangkan"),
        "formal":       ("resmi, sopan, hormat",                   "Gunakan bahasa Indonesia baku"),
    },
    "Malay": {
        "happy":        ("gembira, bahagia, seronok, ceria",      "Ayat pendek dan bertenaga"),
        "sad":          ("sedih, dukacita, pilu, nestapa",         "Ayat perlahan dan penuh perasaan"),
        "romantic":     ("cinta, kasih, romantik, hati",          "Bahasa lembut dan puitis"),
        "angry":        ("marah, murka, geram, berang",            "Ayat tajam dan terus"),
        "calm":         ("tenang, damai, aman",                    "Irama perlahan dan menyenangkan"),
    },
    "Swahili": {
        "happy":        ("furaha, shangwe, raha, fahari",          "Sentensi fupi na zenye nguvu"),
        "sad":          ("huzuni, masikitiko, majonzi",            "Sentensi polepole na za kihisia"),
        "romantic":     ("upendo, mapenzi, pendo, moyo",          "Lugha laini na ya kisanaa"),
        "angry":        ("hasira, ghadhabu, uchungu",              "Sentensi kali na za moja kwa moja"),
        "calm":         ("amani, utulivu, starehe",                "Mdundo wa polepole na wa kupumzika"),
    },
    "French": {
        "happy":        ("bonheur, joie, allégresse, ravissement", "Phrases courtes et énergiques"),
        "sad":          ("tristesse, mélancolie, chagrin, peine",  "Phrases lentes et émouvantes"),
        "romantic":     ("amour, romantisme, tendresse, cœur",     "Langage doux et poétique"),
        "angry":        ("colère, fureur, rage, courroux",         "Phrases vives et directes"),
        "poetic":       ("poésie, métaphore, image, rythme",       "Style baudelairien ou verlainien"),
        "formal":       ("formel, respectueux, soutenu",           "Utiliser le français soutenu"),
        "nostalgic":    ("nostalgie, souvenirs, passé, mélancolie","Écrire avec douce nostalgie"),
    },
    "Spanish": {
        "happy":        ("felicidad, alegría, júbilo, contento",   "Frases cortas y enérgicas"),
        "sad":          ("tristeza, melancolía, pena, dolor",      "Frases lentas y emotivas"),
        "romantic":     ("amor, romance, ternura, corazón",        "Lenguaje suave y poético"),
        "angry":        ("ira, enojo, furia, rabia",               "Frases directas y fuertes"),
        "poetic":       ("poesía, metáfora, imagen, ritmo",        "Al estilo de Neruda o García Lorca"),
        "formal":       ("formal, respetuoso, profesional",        "Usar español culto y formal"),
        "nostalgic":    ("nostalgia, recuerdos, pasado, añoranza", "Escribir con cálida nostalgia"),
    },
    "Portuguese": {
        "happy":        ("felicidade, alegria, júbilo, contentamento","Frases curtas e energéticas"),
        "sad":          ("tristeza, melancolia, saudade, dor",     "Frases lentas e emotivas"),
        "romantic":     ("amor, romance, ternura, coração",        "Linguagem suave e poética"),
        "angry":        ("raiva, ira, fúria, cólera",              "Frases diretas e fortes"),
        "poetic":       ("poesia, metáfora, imagem, ritmo",        "Ao estilo de Fernando Pessoa"),
        "nostalgic":    ("saudade, memórias, passado, nostalgia",  "Escrever com saudade"),
        "formal":       ("formal, respeitoso, profissional",       "Usar português formal"),
    },
    "German": {
        "happy":        ("Glück, Freude, Fröhlichkeit, Heiterkeit","Kurze energische Sätze"),
        "sad":          ("Traurigkeit, Melancholie, Kummer, Leid", "Langsame, gefühlvolle Sätze"),
        "romantic":     ("Liebe, Romantik, Zärtlichkeit, Herz",   "Sanfte, poetische Sprache"),
        "angry":        ("Wut, Ärger, Zorn, Empörung",            "Scharfe, direkte Sätze"),
        "poetic":       ("Poesie, Metapher, Bild, Rhythmus",       "Im Stil von Goethe oder Rilke"),
        "formal":       ("formell, respektvoll, professionell",    "Hochdeutsch und förmlich"),
        "nostalgic":    ("Nostalgie, Erinnerungen, Sehnsucht",    "Mit warmer Nostalgie schreiben"),
    },
    "Italian": {
        "happy":        ("felicità, gioia, allegria, letizia",     "Frasi brevi ed energiche"),
        "sad":          ("tristezza, malinconia, dolore, pena",    "Frasi lente ed emozionanti"),
        "romantic":     ("amore, romanticismo, tenerezza, cuore", "Linguaggio dolce e poetico"),
        "angry":        ("rabbia, ira, furore, indignazione",      "Frasi dirette e forti"),
        "poetic":       ("poesia, metafora, immagine, ritmo",      "In stile dantesco o leopardiano"),
        "formal":       ("formale, rispettoso, professionale",     "Italiano corretto e formale"),
        "nostalgic":    ("nostalgia, ricordi, passato, rimpianto", "Scrivere con dolce nostalgia"),
    },
    "Thai": {
        "happy":        ("ความสุข, ยินดี, สุขใจ, ดีใจ",          "ประโยคสั้น มีพลัง"),
        "sad":          ("เศร้า, โศกเศร้า, เสียใจ, ทุกข์",        "ประโยคช้า เต็มไปด้วยอารมณ์"),
        "romantic":     ("ความรัก, โรแมนติก, อบอุ่น",             "ภาษาอ่อนโยน บทกวี"),
        "angry":        ("โกรธ, โมโห, ฉุนเฉียว, เกรี้ยวกราด",   "ประโยคคม ตรงไปตรงมา"),
        "calm":         ("สงบ, เงียบสงบ, ผ่อนคลาย",              "จังหวะช้า ผ่อนคลาย"),
        "formal":       ("เป็นทางการ, สุภาพ, ราชาศัพท์",          "ใช้ภาษาทางการของไทย"),
    },
    "Burmese": {
        "happy":        ("ပျော်ရွှင်ခြင်း, ဝမ်းမြောက်ခြင်း",     "တိုတောင်းသောဝါကျများ"),
        "sad":          ("ဝမ်းနည်းခြင်း, စိတ်ပျက်ခြင်း",          "နှေးကွေးသောဝါကျများ"),
        "romantic":     ("ချစ်ခြင်းမေတ္တာ, နှလုံးသား",            "နူးညံ့သောဘာသာစကား"),
        "calm":         ("ငြိမ်သက်ခြင်း, ချမ်းသာ",               "နှေးကွေးသောပြေလည်သောဂီတ"),
    },
    "Khmer": {
        "happy":        ("សប្បាយ, រីករាយ, ទូទឹកចិត្ត",           "ប្រយោគខ្លី ស្រស់ស្រាយ"),
        "sad":          ("悲, ទុក្ខ, ព្រួយ",                      "ប្រយោគយឺត ពោរពេញអារម្មណ៍"),
        "calm":         ("សន្តិភាព, ស្ងៀមស្ងាត់",                "ចង្វាក់យឺត រីករាយ"),
    },
    "Georgian": {
        "happy":        ("ბედნიერება, სიხარული, მხიარულება",       "მოკლე ენერგიული წინადადებები"),
        "sad":          ("სევდა, მწუხარება, ტანჯვა",              "ნელი, ემოციური წინადადებები"),
        "romantic":     ("სიყვარული, რომანტიკა, სიტკბო",          "რბილი, პოეტური ენა"),
        "calm":         ("სიმშვიდე, მყუდროება, სიჩუმე",          "ნელი, სასიამოვნო რიტმი"),
    },
    "Armenian": {
        "happy":        ("երջանկություն, ուրախություն, զվարճանք",  "Կարճ, եռանդուն նախադասություններ"),
        "sad":          ("տխրություն, վիշտ, ցավ",                  "Դանդաղ, հուզական նախադասություններ"),
        "romantic":     ("սեր, ռոմանտիկա, քնքշություն",           "Մեղմ, բանաստեղծական լեզու"),
        "calm":         ("խաղաղություն, հանգստություն",            "Դանդաղ, հաճելի ռիթմ"),
    },
}

def get_lang_tone_seed(output_lang, tone):
    """Get vocabulary seed for a specific language+tone combo."""
    if not output_lang:
        return ""
    lang_data = LANG_TONE_VOCAB.get(output_lang, {})
    if not lang_data:
        return ""
    # Exact tone match
    if tone in lang_data:
        words, style_note = lang_data[tone]
        return (f"VOCABULARY SEED for {output_lang} '{tone}' tone:\n"
                f"  Emotional words: {words}\n"
                f"  Style note: {style_note}")
    # Fuzzy match
    tone_lower = tone.lower()
    for key, (words, style_note) in lang_data.items():
        if key in tone_lower or tone_lower in key:
            return (f"VOCABULARY SEED for {output_lang} '{tone}' tone:\n"
                    f"  Emotional words: {words}\n"
                    f"  Style note: {style_note}")
    return ""






# ================== CSS ==================
custom_css = """
.gradio-container {
    min-height: 100vh;
    background: linear-gradient(135deg, #e6f2ff, #f2f9ff) !important;
    font-family: "Segoe UI", Roboto, Arial, sans-serif;
    padding: 30px;
}
h1, h2, h3 { text-align: center; color: #003366; }
label { font-weight: 600; color: #003366; }
input, textarea {
    border-radius: 12px !important;
    border: 1px solid #90c2ff !important;
    padding: 10px !important;
    font-size: 14px !important;
}
input:focus, textarea:focus {
    outline: none !important;
    border-color: #007bff !important;
    box-shadow: 0 0 0 2px rgba(0,123,255,0.2) !important;
}
button.primary {
    background: linear-gradient(135deg, #4da3ff, #007bff) !important;
    color: white !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    font-size: 15px !important;
    padding: 12px 20px !important;
    border: none !important;
    transition: all 0.2s ease-in-out !important;
}
button.primary:hover {
    transform: translateY(-1px) !important;
    background: linear-gradient(135deg, #007bff, #0056b3) !important;
}
.upload-container, .file-preview, [data-testid="file-upload"],
.gr-file-upload, div[class*="upload"] {
    background: #ffffff !important;
    border: 2px dashed #b0cff7 !important;
    border-radius: 14px !important;
    box-shadow: none !important;
}
div[class*="upload"]:hover {
    border-color: #007bff !important;
    background: #f8fbff !important;
}
.gr-dropdown, [data-testid="dropdown"] {
    border-radius: 12px !important;
    border: 1px solid #90c2ff !important;
    background: #ffffff !important;
}
.gr-dropdown:focus-within, [data-testid="dropdown"]:focus-within {
    border-color: #007bff !important;
    box-shadow: 0 0 0 2px rgba(0,123,255,0.2) !important;
}
ul[data-testid="dropdown-options"], .dropdown-options {
    border-radius: 12px !important;
    box-shadow: 0 8px 24px rgba(0,100,255,0.15) !important;
    border: 1px solid #b0cff7 !important;
    max-height: 240px !important;
    overflow-y: auto !important;
}
ul[data-testid="dropdown-options"] li:hover {
    background: #e8f3ff !important;
    color: #003366 !important;
}
.rate-limit-box {
    background: #ffffff;
    border-radius: 16px;
    padding: 18px 28px;
    box-shadow: 0 4px 15px rgba(0,123,255,0.1);
    margin-bottom: 18px;
    display: flex;
    gap: 40px;
}
.rate-item { flex: 1; }
.rate-label { font-size: 11px; font-weight: 700; color: #555; letter-spacing: 1.2px; margin-bottom: 8px; }
.rate-bar-bg { background: #e8f0fe; border-radius: 999px; height: 8px; width: 100%; overflow: hidden; margin-bottom: 6px; }
.rate-bar-fill { height: 100%; border-radius: 999px; transition: width 0.4s ease; }
.rate-text { font-size: 13px; color: #333; }
"""

# ================== RATE DISPLAY ==================
def get_rate_status():
    with rate_limit_lock:
        now = time.time()
        while minute_timestamps and now - minute_timestamps[0] > MINUTE_WINDOW:
            minute_timestamps.popleft()
        while day_timestamps and now - day_timestamps[0] > DAY_WINDOW:
            day_timestamps.popleft()
        return len(minute_timestamps), len(day_timestamps)

def record_one_request():
    with rate_limit_lock:
        now = time.time()
        minute_timestamps.append(now)
        day_timestamps.append(now)

def build_rate_html(min_used, day_used):
    min_pct  = min(100, round(min_used / MAX_REQUESTS_PER_MINUTE * 100))
    day_pct  = min(100, round(day_used / MAX_REQUESTS_PER_DAY * 100))
    min_left = max(0, MAX_REQUESTS_PER_MINUTE - min_used)
    day_left = max(0, MAX_REQUESTS_PER_DAY - day_used)
    mc = "#ef4444" if min_pct >= 80 else "#22c55e"
    dc = "#ef4444" if day_pct >= 80 else "#22c55e"
    return f"""
    <div class="rate-limit-box">
        <div class="rate-item">
            <div class="rate-label">⏱ PER MINUTE</div>
            <div class="rate-bar-bg"><div class="rate-bar-fill" style="width:{min_pct}%;background:{mc};"></div></div>
            <div class="rate-text"><strong>{min_used}/{MAX_REQUESTS_PER_MINUTE}</strong> used &nbsp;·&nbsp; {min_left} left</div>
        </div>
        <div class="rate-item">
            <div class="rate-label">📅 TODAY</div>
            <div class="rate-bar-bg"><div class="rate-bar-fill" style="width:{day_pct}%;background:{dc};"></div></div>
            <div class="rate-text"><strong>{day_used}/{MAX_REQUESTS_PER_DAY}</strong> used &nbsp;·&nbsp; {day_left} left</div>
        </div>
    </div>"""


# ================== DOCUMENT STRUCTURE ==================
H_TAG = "##HEADING##"

def _mark_heading(text):
    return f"{H_TAG} {text.strip()}"

def _is_heading_line(line):
    return line.strip().startswith(H_TAG)

def _heading_text(line):
    return line.strip()[len(H_TAG):].strip()

def clean_for_display(text):
    """Remove internal heading markers — user sees clean readable text."""
    if not text:
        return text
    lines = []
    for line in text.split('\n'):
        lines.append(_heading_text(line) if _is_heading_line(line) else line)
    return re.sub(r'\n{4,}', '\n\n\n', '\n'.join(lines)).strip()


# ================== FILE READING ==================
def _read_pdf_structured(filepath):
    doc       = fitz.open(filepath)
    all_lines = []
    for page_num, page in enumerate(doc):
        for block in page.get_text("dict", sort=True)["blocks"]:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                txt, sz, bold = "", 0, False
                y_top, y_bot  = line["bbox"][1], line["bbox"][3]
                for span in line["spans"]:
                    t = span["text"].strip()
                    if t:
                        txt  += t + " "
                        sz    = max(sz, round(span["size"], 1))
                        bold  = bold or ("bold" in span["font"].lower()) or bool(span["flags"] & 16)
                txt = txt.strip()
                if txt:
                    all_lines.append({"text": txt, "size": sz, "bold": bold,
                                      "y_top": y_top, "y_bot": y_bot, "page": page_num})
    if not all_lines:
        return ""
    body_size  = Counter(l["size"] for l in all_lines).most_common(1)[0][0]
    avg_line_h = body_size * 1.4
    result, prev = [], None
    for line in all_lines:
        is_h = (line["size"] > body_size * 1.15) or (line["bold"] and line["size"] >= body_size * 0.95)
        new_para = prev and (
            line["page"] != prev["page"] or
            line["y_top"] - prev["y_bot"] > avg_line_h * 0.75
        )
        if is_h:
            result.append(f"\n\n{_mark_heading(line['text'])}\n\n")
        else:
            if new_para and result:
                result.append("\n\n")
            result.append(line["text"] + " ")
        prev = line
    full = "".join(result)
    return re.sub(r'\n{4,}', '\n\n\n', re.sub(r'  +', ' ', full)).strip()


def _read_docx_structured(filepath):
    document = docx.Document(filepath)
    result   = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style or any(h in style for h in ["Title", "Subtitle"]):
            result.append(f"\n\n{_mark_heading(text)}\n\n")
        else:
            is_manual_h = False
            if para.runs:
                try:
                    fsize = para.runs[0].font.size.pt if para.runs[0].font.size else None
                except Exception:
                    fsize = None
                if para.runs[0].bold and fsize and fsize > 13:
                    is_manual_h = True
                elif all(r.bold for r in para.runs if r.text.strip()):
                    is_manual_h = True
            result.append(f"\n\n{_mark_heading(text)}\n\n" if is_manual_h else f"\n\n{text}")
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                result.append("\n\n" + " | ".join(cells))
    return re.sub(r'\n{4,}', '\n\n\n', "".join(result)).strip()


def _read_txt_structured(filepath):
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()
    raw    = raw.replace('\r\n', '\n').replace('\r', '\n')
    blocks = re.split(r'\n\s*\n', raw)
    result = []
    for block in blocks:
        lines = [l.rstrip() for l in block.split('\n') if l.strip()]
        if not lines:
            continue
        if len(lines) == 1:
            line = lines[0].strip()
            if ((line.isupper() and len(line) > 2) or
                (line.endswith(':') and len(line.split()) <= 8) or
                (len(line.split()) <= 6 and not line.endswith(('.', ',')) and len(line) > 3)):
                result.append(f"\n\n{_mark_heading(line)}\n\n")
                continue
        merged = " ".join(l.strip() for l in lines)
        result.append(f"\n\n{merged.strip()}")
    return re.sub(r'\n{4,}', '\n\n\n', "".join(result)).strip()


def extract_text_from_file(filepath):
    if not filepath:
        return ""
    try:
        lower = str(filepath).lower()
        if lower.endswith(".pdf"):
            return _read_pdf_structured(filepath)
        elif lower.endswith(".docx"):
            return _read_docx_structured(filepath)
        elif lower.endswith(".txt"):
            return _read_txt_structured(filepath)
        return "⚠️ Unsupported file type. Please upload TXT, PDF, or DOCX."
    except Exception as e:
        return f"⚠️ Error reading file: {e}"

def on_file_upload(filepath):
    raw = extract_text_from_file(filepath)
    return clean_for_display(raw), raw


# ================== CHUNKING ==================
def split_into_chunks(text, chunk_words=CHUNK_SIZE_WORDS):
    paragraphs = [p.strip() for p in re.split(r'\n\n+', text) if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip()]
    chunks, current, wc = [], [], 0
    for para in paragraphs:
        is_h  = _is_heading_line(para)
        words = para.split()
        pw    = len(words) if words else max(1, len(para) // 5)
        if is_h:
            if current:
                chunks.append('\n\n'.join(current))
            current, wc = [para], pw
        elif pw > chunk_words:
            if current:
                chunks.append('\n\n'.join(current))
                current, wc = [], 0
            if words:
                for i in range(0, len(words), chunk_words):
                    chunks.append(' '.join(words[i:i + chunk_words]))
            else:
                step = chunk_words * 5
                for i in range(0, len(para), step):
                    chunks.append(para[i:i + step])
        elif wc + pw > chunk_words:
            chunks.append('\n\n'.join(current))
            current, wc = [para], pw
        else:
            current.append(para)
            wc += pw
    if current:
        chunks.append('\n\n'.join(current))
    return chunks or [text]


# ================== TONE ENGINE ==================
# 60+ tones with synonyms map for maximum coverage
# Format: "canonical_key": (description, writing_instructions)

TONE_PROFILES = {
    # ── Happiness / positivity ─────────────────────────────────────────────
    "happy": (
        "joyful, upbeat, and warm",
        "Radiate genuine happiness. Use bright energetic vocabulary. "
        "Mix short punchy sentences with flowing ones. Let positivity shine naturally."
    ),
    "joyful": (
        "deeply happy, celebratory, and full of life",
        "Express pure delight and celebration. Expansive generous language. "
        "Every sentence should feel like a smile."
    ),
    "excited": (
        "enthusiastic, energetic, and bursting with anticipation",
        "Fast-paced sentences. High energy throughout. "
        "Express wonder, eagerness, and momentum. Build excitement progressively."
    ),
    "cheerful": (
        "bright, lively, and warmly optimistic",
        "Light sunny tone. Short bright sentences. "
        "Friendly and warm — like someone in a genuinely good mood."
    ),
    "hopeful": (
        "optimistic, forward-looking, and full of possibility",
        "Language that points toward a better future. Express belief and possibility. "
        "Even challenges are framed as opportunities."
    ),
    "upbeat": (
        "positive, energetic, and encouraging",
        "Brisk energetic sentences. Positive framing of every idea. "
        "Forward momentum in the rhythm."
    ),
    "optimistic": (
        "positive, confident about the future",
        "Frame everything in terms of what is possible and good. "
        "Bright vocabulary, forward-looking language."
    ),
    "enthusiastic": (
        "eager, passionate, and full of energy",
        "Express strong interest and excitement. High-energy vocabulary. "
        "The writing should feel alive and engaged."
    ),
    "playful": (
        "fun, light-hearted, and full of wit",
        "Use wordplay, lightness, and a sense of fun. "
        "Don't take things too seriously. Warm and inclusive humor."
    ),
    "energetic": (
        "dynamic, vigorous, and full of drive",
        "Short punchy sentences. Action verbs. Fast rhythm. "
        "The text should feel like it's moving."
    ),

    # ── Love / warmth ──────────────────────────────────────────────────────
    "romantic": (
        "tender, intimate, and deeply affectionate",
        "Soft flowing sentences with poetic imagery. "
        "Express longing, warmth, and devotion. Sensory words — touch, warmth, closeness."
    ),
    "loving": (
        "deeply affectionate, nurturing, and unconditionally warm",
        "Express genuine deep love. Protective caring language. "
        "Total acceptance and warmth. Every sentence should feel like an embrace."
    ),
    "tender": (
        "gentle, caring, and softly emotional",
        "Very soft and gentle language. Careful word choices. "
        "Express care and affection without intensity — quiet warmth."
    ),
    "affectionate": (
        "warm, fond, and full of care",
        "Express genuine fondness and care. Warm familiar language. "
        "Sentences that feel like a gentle touch."
    ),
    "passionate": (
        "intensely emotional, fervent, and deeply felt",
        "Strong vivid language. Express deep feeling with intensity. "
        "Everything matters. Full emotional commitment in every sentence."
    ),

    # ── Sadness / pain ────────────────────────────────────────────────────
    "sad": (
        "melancholic, sorrowful, and heavy-hearted",
        "Slow reflective sentences. Words that carry weight and longing. "
        "Express grief or quiet despair. Let sadness breathe without rushing to resolution."
    ),
    "melancholic": (
        "quietly sad, wistful, and reflective",
        "A gentle persistent sadness that colors everything. "
        "Slow measured sentences. Bittersweet imagery."
    ),
    "heartbroken": (
        "devastated, raw, and deeply wounded",
        "Intense emotional pain. Fragmented thoughts, overwhelmed feeling. "
        "Allow the rawness without rushing toward resolution."
    ),
    "gloomy": (
        "dark, heavy, and without much hope",
        "Heavy slow language. Everything feels weighted. "
        "The world is grey. No brightness or uplift."
    ),
    "sorrowful": (
        "deeply grieving and mournful",
        "Deep grief expressed with dignity. "
        "Long reflective sentences. Quiet and serious."
    ),
    "lonely": (
        "isolated, longing, and quietly aching",
        "Express the weight of absence. Introspective sentences that turn inward. "
        "Longing for connection without bitterness."
    ),
    "disappointed": (
        "let down, deflated, and quietly sorrowful",
        "Unmet expectations expressed with restrained sadness. Not angry — just heavy. "
        "Show that hope existed before it was lost."
    ),
    "regretful": (
        "remorseful, reflective, and tinged with sorrow",
        "Look backward with sadness at choices made. Past tense. "
        "Express a genuine wish things had been different."
    ),
    "desperate": (
        "urgent, overwhelmed, and emotionally raw",
        "Short fragmented sentences. High urgency. "
        "Express a sense of running out of options. Rawness and vulnerability."
    ),

    # ── Anger / intensity ─────────────────────────────────────────────────
    "angry": (
        "fierce, intense, and emotionally charged",
        "Short sharp sentences. Strong forceful verbs. "
        "Express frustration or outrage. Rhetorical questions and exclamations. No softness."
    ),
    "aggressive": (
        "confrontational, bold, and unapologetic",
        "Direct blunt sentences. Commanding verbs. "
        "Express challenge and dominance. Zero hedging."
    ),
    "frustrated": (
        "irritated, fed up, and strained",
        "Express the exhaustion of dealing with obstacles. "
        "Short sentences, some rhetorical questions. "
        "Not quite rage but clearly at the end of patience."
    ),
    "fierce": (
        "powerful, intense, and unrelenting",
        "Strong bold language. Every sentence has force. "
        "Uncompromising and direct."
    ),
    "cold": (
        "detached, icy, and emotionally distant",
        "Flat unemotional language. Short precise sentences. "
        "No warmth whatsoever. Deliberate emotional distance."
    ),

    # ── Fear / tension ────────────────────────────────────────────────────
    "fearful": (
        "anxious, tense, and on edge",
        "Fragmented sentences and uncertain language. "
        "Express dread and unease. Short choppy sentences build tension."
    ),
    "anxious": (
        "worried, nervous, and unsettled",
        "Restless language. Sentences that circle back. "
        "Express constant low-level worry. Uncertainty in word choices."
    ),
    "dark": (
        "sinister, shadowy, and unsettling",
        "Use dark imagery. Shadows, decay, foreboding. "
        "Unsettling atmosphere in every sentence."
    ),
    "mysterious": (
        "enigmatic, suspenseful, and intriguing",
        "Leave things unsaid — imply more than you state. "
        "Short cryptic statements alongside longer atmospheric ones."
    ),
    "tense": (
        "high-stakes, pressured, and suspenseful",
        "Short sentences. Every word tight. "
        "No relaxation — constant forward pressure."
    ),

    # ── Complex / mixed ───────────────────────────────────────────────────
    "nostalgic": (
        "wistful, warm, and reflective of the past",
        "Memory-laden language. Evoke sensory details from earlier times. "
        "Bittersweet — happy memories tinged with the passage of time."
    ),
    "bittersweet": (
        "simultaneously joyful and sorrowful",
        "Hold joy and sadness in the same sentence. "
        "Express that something beautiful is also passing or lost."
    ),
    "sarcastic": (
        "dry, ironic, and witty with a sharp edge",
        "Say the opposite of what you mean with exaggerated politeness. "
        "Understatement and deadpan tone. Humor from restraint."
    ),
    "ironic": (
        "deliberately contrary, wryly observational",
        "Present situations highlighting contradiction. "
        "Understated and detached. Let reality speak for itself."
    ),
    "humorous": (
        "playful, funny, and light-hearted",
        "Wit, wordplay, and unexpected comparisons. "
        "Set up and subvert expectations. Don't explain the joke."
    ),
    "dramatic": (
        "intense, theatrical, and emotionally heightened",
        "Amplify every emotion. Use contrast — silence vs chaos, love vs loss. "
        "Rhetorical questions, repetition for effect. Everything matters deeply."
    ),
    "surprised": (
        "astonished, caught off guard, and wide-eyed",
        "Express genuine surprise and disbelief. "
        "Short exclamatory sentences. The unexpected is everywhere."
    ),
    "bored": (
        "disengaged, flat, and going through the motions",
        "Flat monotonous language. Long tedious sentences. "
        "Nothing is interesting. Minimal enthusiasm."
    ),

    # ── Warmth / social ───────────────────────────────────────────────────
    "empathetic": (
        "warm, understanding, and compassionate",
        "Acknowledge feelings before perspective. Gentle inclusive language. "
        "Validate the reader's experience. Sentences like a warm hand on the shoulder."
    ),
    "compassionate": (
        "deeply caring, gentle, and moved by suffering",
        "Genuine concern. Soft gentle language. "
        "Truly present with the person's experience."
    ),
    "grateful": (
        "warm, appreciative, and deeply thankful",
        "Express genuine thankfulness. Heartfelt sincere language. "
        "Name what is valued specifically."
    ),
    "gentle": (
        "soft, careful, and considerate",
        "Very soft language with no sharp edges. "
        "Careful considerate word choices. Nothing aggressive or sudden."
    ),

    # ── Strength / power ──────────────────────────────────────────────────
    "motivational": (
        "inspiring, empowering, and action-driven",
        "Strong action verbs. Address the reader directly. "
        "Build momentum with rhythm. Express unshakeable belief in the reader's potential."
    ),
    "confident": (
        "assured, decisive, and self-possessed",
        "Declarative sentences. No hedging. State things as facts. "
        "Active voice. Express certainty and clarity."
    ),
    "proud": (
        "dignified, accomplished, and deeply satisfied",
        "Express earned achievement. Measured confident language. "
        "Acknowledge effort and meaning behind accomplishment."
    ),
    "authoritative": (
        "commanding, expert, and definitively stated",
        "Speak as an unquestioned expert. Declarative statements. "
        "Others look to you for the final word."
    ),
    "bold": (
        "daring, direct, and unapologetically strong",
        "Strong direct language. No hedging or qualification. "
        "Express ideas with force and certainty."
    ),
    "inspiring": (
        "uplifting, visionary, and deeply moving",
        "Speak to what is possible and meaningful. "
        "Language that stirs the soul. Connect individual to universal."
    ),

    # ── Formal / professional ─────────────────────────────────────────────
    "formal": (
        "professional, precise, and respectful",
        "Complete sentences, proper grammar, formal vocabulary. "
        "No contractions, slang, or emotional language. Authoritative but respectful."
    ),
    "professional": (
        "confident, clear, and authoritative",
        "Precise language and structured sentences. "
        "Active voice, concrete statements, no fluff."
    ),
    "academic": (
        "scholarly, analytical, and evidence-based",
        "Precise technical vocabulary. Build arguments logically. "
        "Objective and analytical — emotion stays out."
    ),
    "diplomatic": (
        "tactful, balanced, and carefully worded",
        "Present all sides fairly. Avoid polarizing language. "
        "Express firmness gently."
    ),
    "neutral": (
        "balanced, impartial, and without emotional bias",
        "Completely balanced language. No emotional loading. "
        "Just the facts, presented fairly from all sides."
    ),
    "pessimistic": (
        "negative, doubtful, and expecting the worst",
        "Frame everything in terms of problems and failures. "
        "Doubt everything. Nothing is quite good enough."
    ),
    "sincere": (
        "genuine, heartfelt, and completely authentic",
        "Strip away performance. Speak directly from the heart. "
        "Simple honest language. Every word should feel true."
    ),

    # ── Literary / creative ───────────────────────────────────────────────
    "poetic": (
        "lyrical, metaphorical, and rhythmically rich",
        "Vivid imagery, metaphors, and sensory language. "
        "Vary sentence length for rhythm. Elevate the mundane into something beautiful."
    ),
    "descriptive": (
        "richly detailed, vivid, and sensory",
        "Paint pictures with words. Use all five senses. "
        "Specific concrete details. Make the reader see and feel what you describe."
    ),
    "persuasive": (
        "compelling, logical, and emotionally resonant",
        "Clear logical arguments. Address objections preemptively. "
        "Evidence and emotional appeal together."
    ),
    "storytelling": (
        "engaging, narrative, and immersive",
        "Draw the reader in as if telling a story. "
        "Scene-setting, tension, and forward pull. Show don't tell."
    ),

    # ── Calm / reflective ─────────────────────────────────────────────────
    "calm": (
        "peaceful, measured, and serene",
        "Long flowing sentences with gentle rhythm. Soft soothing language. "
        "No urgency or tension. Express tranquility and stillness."
    ),
    "peaceful": (
        "still, harmonious, and gently warm",
        "Very soft gentle language. No conflict or tension anywhere. "
        "A sense of everything being right."
    ),
    "reflective": (
        "thoughtful, introspective, and quietly wise",
        "Slow measured language. Turn inward to explore meaning. "
        "Questions are as important as answers."
    ),
    "philosophical": (
        "contemplative, deep, and exploring fundamental questions",
        "Engage with ideas beyond the surface. Use questions to open thought. "
        "Patient language that explores without needing resolution."
    ),
    "serious": (
        "grave, earnest, and deeply sincere",
        "No lightness. Every word carries full weight. "
        "Express matters of deep importance with matching gravity."
    ),
    "meditative": (
        "deeply still, present, and awareness-centered",
        "Ultra-slow deliberate sentences. Present moment awareness. "
        "Each word placed with intention."
    ),
}

# Synonym map — maps common user inputs that are NOT in TONE_PROFILES
# to the closest canonical profile key
TONE_SYNONYMS = {
    # happy family
    "delighted":   "joyful",      "ecstatic":    "excited",
    "gleeful":     "happy",       "merry":       "cheerful",
    "buoyant":     "upbeat",      "radiant":     "joyful",
    "vibrant":     "energetic",   "lively":      "cheerful",
    "thrilled":    "excited",     "elated":      "joyful",
    "jubilant":    "joyful",      "content":     "calm",
    "satisfied":   "proud",       "blissful":    "peaceful",
    # sad family
    "depressed":   "sad",         "miserable":   "sad",
    "grief":       "sorrowful",   "grieving":    "sorrowful",
    "mourning":    "sorrowful",   "dejected":    "disappointed",
    "despondent":  "gloomy",      "hopeless":    "gloomy",
    "forlorn":     "lonely",      "wistful":     "nostalgic",
    "anguished":   "heartbroken", "devastated":  "heartbroken",
    "hurt":        "heartbroken", "broken":      "heartbroken",
    "tearful":     "sad",         "weeping":     "sad",
    # angry family
    "furious":     "angry",       "enraged":     "angry",
    "irritated":   "frustrated",  "annoyed":     "frustrated",
    "bitter":      "angry",       "hostile":     "aggressive",
    "resentful":   "frustrated",  "outraged":    "angry",
    "wrathful":    "angry",       "indignant":   "angry",
    # fear family
    "scared":      "fearful",     "terrified":   "fearful",
    "nervous":     "anxious",     "worried":     "anxious",
    "uneasy":      "anxious",     "paranoid":    "anxious",
    "dread":       "fearful",     "horrified":   "fearful",
    "panicked":    "fearful",     "unsettled":   "anxious",
    # love family
    "adoring":     "loving",      "devoted":     "romantic",
    "warm":        "loving",      "caring":      "compassionate",
    "nurturing":   "loving",      "intimate":    "romantic",
    "sensual":     "romantic",    "longing":     "romantic",
    # calm family
    "relaxed":     "calm",        "tranquil":    "peaceful",
    "serene":      "peaceful",    "composed":    "calm",
    "gentle":      "gentle",      "soft":        "gentle",
    "quiet":       "calm",        "still":       "meditative",
    # strength family
    "powerful":    "bold",        "strong":      "confident",
    "determined":  "confident",   "resolute":    "confident",
    "assertive":   "confident",   "decisive":    "confident",
    "courageous":  "bold",        "brave":       "bold",
    # misc
    "witty":       "humorous",    "funny":       "humorous",
    "comical":     "humorous",    "dry":         "sarcastic",
    "wry":         "sarcastic",   "sardonic":    "sarcastic",
    "cynical":     "sarcastic",   "lyrical":     "poetic",
    "vivid":       "descriptive", "narrative":   "storytelling",
    "thoughtful":  "reflective",  "pensive":     "reflective",
    "wise":        "philosophical","deep":       "philosophical",
    "grave":       "serious",     "solemn":      "serious",
    "earnest":     "sincere",     "honest":      "sincere",
    "authentic":   "sincere",     "genuine":     "sincere",
    "inspiring":   "inspirational","uplifting":  "inspirational",
    "electric":    "excited",     "stirring":    "dramatic",
    "intense":     "dramatic",    "theatrical":  "dramatic",
}

def _get_tone_guidance(tone_input):
    """
    Comprehensive tone matching:
    1. Exact match in TONE_PROFILES
    2. Synonym map lookup
    3. Partial substring match
    4. Word-by-word scan
    5. Intelligent fallback with custom instructions
    """
    raw   = tone_input.strip()
    lower = raw.lower()

    # 1. Exact match
    if lower in TONE_PROFILES:
        return TONE_PROFILES[lower]

    # 2. Synonym map
    if lower in TONE_SYNONYMS:
        return TONE_PROFILES[TONE_SYNONYMS[lower]]

    # 3. Check each word in synonyms map
    words = re.findall(r'[a-z]+', lower)
    for word in words:
        if word in TONE_SYNONYMS:
            return TONE_PROFILES[TONE_SYNONYMS[word]]

    # 4. Substring match in TONE_PROFILES keys
    for key in TONE_PROFILES:
        if key in lower or lower in key:
            return TONE_PROFILES[key]

    # 5. Word-by-word scan in TONE_PROFILES
    for word in words:
        if word in TONE_PROFILES:
            return TONE_PROFILES[word]
        for key in TONE_PROFILES:
            if word in key:
                return TONE_PROFILES[key]

    # 6. Intelligent fallback — custom instructions for any unknown tone
    return (
        f"deeply and unmistakably {raw}",
        f"Fully embody the '{raw}' tone in every sentence. "
        f"Ask yourself: how would someone genuinely feeling '{raw}' write? "
        f"What vocabulary, sentence length, rhythm, and imagery comes naturally to this state? "
        f"Apply those choices consistently throughout. "
        f"Every reader should unmistakably feel '{raw}' from the text, in any language."
    )


# ================== LANGUAGE INSTRUCTION BUILDER ==================
def _build_lang_instruction(input_lang, output_lang):
    """
    Build crystal-clear language instructions for the model.
    Handles all combinations correctly including same-language tone change.
    """
    il = (input_lang or "").strip()
    ol = (output_lang or "").strip()

    if il and ol and il.lower() != ol.lower():
        # Translation + tone change
        return (
            f"LANGUAGE TASK: The input is written in {il}. "
            f"You MUST write your entire response in {ol} only. "
            f"Every single word must be in {ol}. "
            f"Do not include any {il} words. "
            f"This is a translation AND tone transformation simultaneously. "
            f"Express the '{ol}' language's natural emotional vocabulary for this tone — "
            f"do not just translate English emotional words literally."
        )
    elif ol and (not il or il.lower() == ol.lower()):
        # Same language, just tone change
        return (
            f"LANGUAGE: Write your entire response in {ol}. "
            f"Keep the same language — only the tone and emotional style changes. "
            f"Use {ol}'s natural vocabulary and expressions for this emotional tone."
        )
    elif il and not ol:
        # Keep input language
        return (
            f"LANGUAGE: The text is in {il}. "
            f"Respond entirely in {il}. Do not translate or switch languages. "
            f"Use {il}'s natural vocabulary and expressions for this emotional tone."
        )
    else:
        # Auto-detect
        return (
            "LANGUAGE: Identify the language of the input and write your entire response "
            "in that exact same language. Do not translate. "
            "Use that language's natural vocabulary and expressions for this emotional tone."
        )


# ================== SYSTEM PROMPT ==================
def build_system_prompt(tone, input_lang, output_lang):
    """
    Builds a powerful, language-specific, tone-precise system prompt.
    
    Key upgrades:
    - Injects native vocabulary seeds for 38 languages so model uses
      the language's OWN emotional tradition, not translated English
    - Script-family awareness for proper writing direction/style hints
    - Tone-specific rhythm and sentence structure instructions
    """
    tone_desc, tone_instr = _get_tone_guidance(tone)
    lang_instruction      = _build_lang_instruction(input_lang, output_lang)
    
    # Get vocabulary seed for the output language + tone combo
    vocab_seed = get_lang_tone_seed(output_lang or "", tone)
    
    # Script family hint
    script = get_script_family(output_lang or "")
    script_hints = {
        "indic":   "Write flowing connected script. Avoid breaking natural word boundaries. "
                   "Use the script's natural sandhi and compound word formations.",
        "arabic":  "Write right-to-left. Use Arabic literary tradition. "
                   "Classical rhetorical devices (saj', tawriya) are appropriate for poetic/formal tones.",
        "cjk":     "Use appropriate character density. "
                   "Classical 4-character idioms (成語/四字熟語/사자성어) add depth for formal/poetic tones.",
        "cyrillic":"Use the full range of the language's emotional vocabulary. "
                   "Rich Russian/Slavic emotional tradition.",
        "latin":   "",
        "rtl":     "Write right-to-left script naturally.",
    }
    script_note = script_hints.get(script, "")

    vocab_section = ""
    if vocab_seed:
        vocab_section = f"\n\n━━ NATIVE VOCABULARY GUIDE ━━\n{vocab_seed}"
    if script_note:
        vocab_section += f"\n\nSCRIPT NOTE: {script_note}"

    return (
        f"You are a master writer and emotional stylist, fluent in all world languages "
        f"and their literary traditions.\n"
        f"Your task: rewrite the given text to powerfully express the '{tone}' tone.\n\n"

        f"━━ TARGET TONE ━━\n"
        f"{tone_desc}\n\n"

        f"━━ HOW TO WRITE THIS TONE ━━\n"
        f"{tone_instr}"
        f"{vocab_section}\n\n"

        f"━━ RULES ━━\n"
        f"1. MEANING: Preserve 100% of the original meaning. "
        f"Only the emotional register, vocabulary, and style change — never the facts.\n\n"

        f"2. TONE DEPTH: Every single sentence must breathe the '{tone}' emotion — "
        f"vocabulary, rhythm, sentence length, and punctuation all serve the tone. "
        f"Do not just change a few words.\n\n"

        f"3. {lang_instruction}\n\n"

        f"4. AUTHENTIC EMOTIONAL VOCABULARY: Use words that carry genuine '{tone}' "
        f"emotional weight in the output language's OWN literary and everyday tradition. "
        f"Do NOT translate English emotional phrases literally — find how this emotion "
        f"is actually expressed by native writers of that language.\n\n"

        f"5. SENTENCE RHYTHM: Match sentence rhythm to the emotion:\n"
        f"   • Short sharp sentences → anger, fear, urgency, excitement, aggressive\n"
        f"   • Long flowing sentences → sadness, romance, calm, poetic, reflective, philosophical\n"
        f"   • Varied dynamic rhythm → dramatic, nostalgic, bittersweet\n"
        f"   • Measured even rhythm → formal, professional, sincere\n\n"

        f"6. HEADINGS: If the input contains standalone title lines "
        f"(short, no body text), keep them as clean lines with a blank line after.\n\n"

        f"7. OUTPUT: Return ONLY the rewritten text. "
        f"Zero preamble, explanation, labels, or commentary.\n\n"

        f"8. LENGTH: Keep approximately the same length as the input.\n\n"

        f"9. CHUNK: This may be one section of a larger document. "
        f"Do not add a title, introduction, conclusion, or summary."
    )


# ================== SSE STREAMER ==================
def _sse_tokens(system_prompt, chunk_text):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "HTTP-Referer":  "https://github.com/your-repo",
        "X-Title":       "Emotion Text Style Transfer",
        "Content-Type":  "application/json",
        "Accept":        "text/event-stream",
    }
    payload = {
        "model":   MODEL,
        "stream":  True,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": chunk_text},
        ],
        "temperature": 0.75,
        "max_tokens":  MAX_TOKENS_PER_CHUNK,
    }
    resp = requests.post(API_URL, headers=headers, json=payload,
                         stream=True, timeout=(15, 120))
    if resp.status_code == 429:
        raise RateLimitError(int(resp.headers.get("Retry-After", RATE_WAIT_SECS)))
    if resp.status_code >= 500:
        raise requests.exceptions.ConnectionError(f"Server error {resp.status_code}")
    resp.raise_for_status()
    done = False
    for raw in resp.iter_lines():
        if not raw:
            continue
        line = raw.decode("utf-8") if isinstance(raw, bytes) else raw
        if not line.startswith("data:"):
            continue
        data = line[5:].strip()
        if data == "[DONE]":
            done = True
            break
        try:
            obj = json.loads(data)
        except Exception:
            continue
        if "error" in obj:
            err_msg  = obj["error"].get("message", str(obj["error"]))
            err_code = obj["error"].get("code", 0)
            if err_code == 429 or "rate" in err_msg.lower():
                raise RateLimitError(RATE_WAIT_SECS)
            raise RuntimeError(f"API error: {err_msg}")
        choices = obj.get("choices")
        if not choices:
            continue
        token = choices[0].get("delta", {}).get("content", "")
        if token:
            yield token
    if not done:
        raise requests.exceptions.ChunkedEncodingError("Stream closed before [DONE]")


class RateLimitError(Exception):
    def __init__(self, wait_secs):
        self.wait_secs = wait_secs
        super().__init__(f"Rate limit — wait {wait_secs}s")


# ================== REPETITION GUARD ==================
def _is_repeating(text):
    if len(text) < 60:
        return False
    tokens = text.split()
    if len(tokens) >= 20 and len(set(tokens[-20:])) <= 2:
        return True
    tail = text[-200:]
    ngrams = [tail[i:i+10] for i in range(0, len(tail)-10, 10)]
    if len(ngrams) >= 8:
        top = max(set(ngrams), key=ngrams.count)
        if ngrams.count(top) / len(ngrams) >= 0.75:
            return True
    tail2 = text[-300:]
    for plen in range(2, 41):
        pat = tail2[-plen:]
        if not pat.strip():
            continue
        count, pos = 0, len(tail2) - plen
        while pos >= plen and tail2[pos-plen:pos] == pat:
            count += 1
            pos -= plen
        if count >= 12:
            return True
    return False


# ================== CHUNK PREP FOR API ==================
def _chunk_for_api(chunk):
    """Strip internal heading markers before sending to model."""
    lines = []
    for line in chunk.split('\n'):
        lines.append(_heading_text(line) if _is_heading_line(line) else line)
    return '\n'.join(lines).strip()


# ================== MAIN GENERATOR ==================
def generate_tone_variation(text, tone, input_lang, output_lang):
    mu, du = get_rate_status()
    if not text or not text.strip():
        yield "⚠️ Please enter some text or upload a file first.", build_rate_html(mu, du)
        return
    if not tone or not tone.strip():
        yield "⚠️ Please specify an emotion or tone.", build_rate_html(mu, du)
        return

    text        = text.strip()
    tone        = tone.strip()
    input_lang  = str(input_lang).strip() if input_lang else ""
    output_lang = str(output_lang).strip() if output_lang else ""

    system_prompt = build_system_prompt(tone, input_lang, output_lang)
    chunks        = split_into_chunks(text, CHUNK_SIZE_WORDS)
    n             = len(chunks)
    output        = ""

    for idx, chunk in enumerate(chunks):
        if idx > 0 and idx % CHUNKS_PER_BATCH == 0:
            for remaining in range(RATE_WAIT_SECS, 0, -1):
                mu2, du2 = get_rate_status()
                yield (output + f"\n\n⏳ {idx}/{n} sections done — waiting {remaining}s…"),\
                      build_rate_html(mu2, du2)
                time.sleep(1)

        record_one_request()
        attempt      = 0
        chunk_output = ""
        api_chunk    = _chunk_for_api(chunk)

        while True:
            try:
                rep_hit = False
                for token in _sse_tokens(system_prompt, api_chunk):
                    chunk_output += token
                    if _is_repeating(chunk_output):
                        chunk_output = chunk_output.rstrip()
                        chunk_output += "\n\n[⚠️ Repetition detected — please regenerate this section.]"
                        rep_hit = True
                        break
                    mu2, du2 = get_rate_status()
                    yield output + chunk_output, build_rate_html(mu2, du2)
                if rep_hit:
                    mu2, du2 = get_rate_status()
                    yield output + chunk_output, build_rate_html(mu2, du2)
                break

            except RateLimitError as e:
                chunk_output = ""
                for remaining in range(e.wait_secs, 0, -1):
                    mu2, du2 = get_rate_status()
                    yield (output + f"\n\n⏳ Rate limit — resuming in {remaining}s…"),\
                          build_rate_html(mu2, du2)
                    time.sleep(1)

            except (requests.exceptions.ConnectionError,
                    requests.exceptions.ChunkedEncodingError,
                    requests.exceptions.ReadTimeout,
                    requests.exceptions.Timeout,
                    ConnectionResetError, BrokenPipeError, OSError) as e:
                attempt += 1
                if attempt > MAX_RETRIES:
                    chunk_output = f"[⚠️ Section {idx+1} failed after {MAX_RETRIES} retries: {e}]"
                    break
                delay = RETRY_DELAYS[min(attempt-1, len(RETRY_DELAYS)-1)]
                chunk_output = ""
                for remaining in range(delay, 0, -1):
                    mu2, du2 = get_rate_status()
                    yield (output + f"\n\n🔄 Retrying section {idx+1} in {remaining}s "
                           f"(attempt {attempt}/{MAX_RETRIES})…"), build_rate_html(mu2, du2)
                    time.sleep(1)

            except Exception as e:
                chunk_output = f"[⚠️ Section {idx+1} error: {e}]"
                break

        output += ("\n\n" + chunk_output) if output and chunk_output else chunk_output
        mu2, du2 = get_rate_status()
        yield output, build_rate_html(mu2, du2)

    mu2, du2 = get_rate_status()
    yield output, build_rate_html(mu2, du2)


# ================== DOWNLOAD BUILDER ==================
EXTENSION_MAP = {
    "TXT":      ".txt",
    "DOCX":     ".docx",
    "Markdown": ".md",
    "HTML":     ".html",
}

def _safe_filename(name, fmt):
    name = str(name or "transformed_output").strip()
    safe = "".join(c for c in name if c.isalnum() or c in " _-").strip()
    safe = safe.replace(" ", "_")
    safe = safe or "transformed_output"

    ext = EXTENSION_MAP.get(fmt, ".txt")
    if not safe.lower().endswith(ext.lower()):
        safe += ext
    return safe

def _parse_output_blocks(text):
    """Detect headings heuristically from clean AI output."""
    blocks = []
    for block in [b.strip() for b in re.split(r'\n\n+', text) if b.strip()]:
        lines = block.split('\n')
        words = block.split()
        is_h  = (len(lines) == 1 and 1 <= len(words) <= 10 and
                 len(block) < 120 and
                 not block.endswith(('.', ',', '!', '?', ';', ':')))
        blocks.append(("heading" if is_h else "para", block))
    return blocks
def clean_text(text):
    replacements = {
        "—": "-",
        "–": "-",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'"
    }

    for k, v in replacements.items():
        text = text.replace(k, v)

    return text

def build_download_file(text, fmt, filename):
    try:
        t = str(text or "").strip()
        if not t:
            return None, "⚠️ Nothing to download — generate output first."

        fmt = str(fmt or "TXT").strip()
        filename = str(filename or "output").strip()

        safe = _safe_filename(filename, fmt)

        if not isinstance(safe, str) or not safe.strip():
            safe = "output.txt"

        filepath = os.path.join(tempfile.gettempdir(), safe)
        blocks = _parse_output_blocks(t)

        if fmt == "TXT":
            lines = []
            for kind, content in blocks:
                content = str(content or "")
                if kind == "heading":
                    sep = "=" * len(content)
                    lines.append(f"{sep}\n{content}\n{sep}")
                else:
                    lines.append(content)

            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n\n".join(lines))

        elif fmt == "DOCX":
            doc = DocxDocument()
            ts  = doc.add_paragraph()
            ts_run = ts.add_run(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(120, 120, 120)
            doc.add_paragraph()
            for kind, content in blocks:
                if kind == "heading":
                    h = doc.add_heading(content, level=1)
                    if h.runs:
                        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                else:
                    p = doc.add_paragraph(content)
                    p.paragraph_format.space_after = Pt(8)
            doc.save(filepath)
        elif fmt == "Markdown":
            lines = [f"*Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}*", "", "---", ""]
            for kind, content in blocks:
                content = str(content or "")
                lines.append(f"## {content}" if kind == "heading" else content)
                lines.append("")

            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))

        elif fmt == "HTML":
            parts = []
            for kind, content in blocks:
                content = str(content or "")
                tag = "h2" if kind == "heading" else "p"
                parts.append(f"  <{tag}>{html_lib.escape(content)}</{tag}>")

            html_content = (
                f'<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"/>'
                f'<meta name="viewport" content="width=device-width,initial-scale=1.0"/>'
                f'<title>Transformed Text</title><style>'
                f'body{{font-family:"Segoe UI",Roboto,Arial,sans-serif;max-width:820px;'
                f'margin:48px auto;padding:0 24px;background:#f4f8ff;color:#222;line-height:1.85;}}'
                f'h2{{color:#003366;border-bottom:1px solid #90c2ff;padding-bottom:6px;margin-top:32px;}}'
                f'.meta{{font-size:12px;color:#888;margin-bottom:28px;}}'
                f'p{{margin-bottom:16px;font-size:15px;}}'
                f'</style></head><body>'
                f'<div class="meta">Generated: {time.strftime("%Y-%m-%d %H:%M:%S")}</div>'
                f'\n' + '\n'.join(parts) + '\n</body></html>'
            )

            with open(filepath, "w", encoding="utf-8") as f:
                f.write(html_content)

        else:
            return None, f"⚠️ Unsupported format: {fmt}"

        return filepath, f"✅ **{safe}** is ready — click the file below to download."

    except Exception as e:
        return None, f"⚠️ Error creating {fmt} file: {type(e).__name__}: {e}"

# ================== GRADIO UI ==================
with gr.Blocks(title="Emotion Text Style Transfer") as demo:

    gr.Markdown("# 🎨 Emotion Based Text Style Transfer")
    gr.Markdown(
        "<div style='text-align:center;color:#444;margin-bottom:10px;'>"
        "Transform your text into any tone, style, and language using AI.<br>"
        "<small>Unlimited pages · word-by-word streaming · structure-aware reading · "
        "full Unicode download · 60+ tones</small>"
        "</div>"
    )

    rate_display   = gr.HTML(value=build_rate_html(0, 0))
    raw_text_state = gr.State("")

    with gr.Row():
        with gr.Column(scale=1):
            text_input = gr.Textbox(
                label="✏️ Enter your text",
                placeholder="Type or paste your text here...",
                lines=6,
            )
            file_input = gr.File(
                label="📎 Or upload a file (TXT · DOCX · PDF)",
                file_types=[".txt", ".docx", ".pdf"],
                file_count="single",
                type="filepath",
            )
            tone_input = gr.Textbox(
                label="🎭 Emotion / Tone",
                placeholder="e.g., happy, sad, romantic, angry, poetic, nostalgic, desperate...",
            )
            with gr.Row():
                input_lang = gr.Dropdown(
                    choices=SUPPORTED_LANGUAGES, label="🌐 Input Language",
                    value="English", allow_custom_value=True, filterable=True,
                    info="Type to search any language",
                )
                output_lang = gr.Dropdown(
                    choices=SUPPORTED_LANGUAGES, label="🌐 Output Language",
                    value="English", allow_custom_value=True, filterable=True,
                    info="Type to search any language",
                )
            generate_btn = gr.Button("✨ Generate Tone Variation", variant="primary", size="lg")

        with gr.Column(scale=1):
            output = gr.Textbox(
                label="📝 Modified Text",
                placeholder="Your transformed text will appear here...",
                lines=15,
                interactive=False,
            )
            gr.Markdown(
                "<div style='font-weight:700;color:#003366;"
                "font-size:15px;margin-top:16px;margin-bottom:8px;'>"
                "⬇️ Download Output</div>"
            )
            with gr.Row():
                dl_filename = gr.Textbox(
                    label="📝 File name (no extension needed)",
                    value="transformed_output",
                    placeholder="e.g. my_output",
                    scale=3,
                )
                dl_format = gr.Dropdown(
                    choices=list(EXTENSION_MAP.keys()),
                    value="TXT",
                    label="📄 Format",
                    scale=1,
                )
            dl_btn  = gr.Button("⬇️ Download", variant="primary", size="lg")
            dl_file = gr.File(
                label="📥 Your file is ready — click to download",
                visible=False, interactive=False,
            )
            dl_status = gr.Markdown(
                value="<span style='color:#888;font-size:13px;'>"
                      "Type a file name, choose a format, then click Download.</span>"
            )

    # ── Wire events ───────────────────────────────────────────────────────
    file_input.change(
        fn=on_file_upload,
        inputs=[file_input],
        outputs=[text_input, raw_text_state]
    )
    text_input.change(
        fn=lambda _: "",
        inputs=[text_input],
        outputs=[raw_text_state]
    )

    def _run_generate(display_text, raw_text, tone, in_lang, out_lang):
        actual = raw_text.strip() if raw_text and raw_text.strip() else display_text
        yield from generate_tone_variation(actual, tone, in_lang, out_lang)

    generate_btn.click(
        fn=_run_generate,
        inputs=[text_input, raw_text_state, tone_input, input_lang, output_lang],
        outputs=[output, rate_display],
    )

    def _do_download(text, fmt, filename):
        path, msg = build_download_file(text, fmt, filename)
        if path:
            return gr.update(value=path, visible=True), msg
        return gr.update(visible=False), msg

    dl_btn.click(
        fn=_do_download,
        inputs=[output, dl_format, dl_filename],
        outputs=[dl_file, dl_status],
    )
    

if __name__ == "__main__":
    demo.queue()
    demo.launch(share=True, css=custom_css)